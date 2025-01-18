from qfluentwidgets import TextEdit,CommandBar,Action
from PyQt6.QtWidgets import QWidget, QVBoxLayout
from PyQt6.QtGui import QFont, QTextListFormat
from PyQt6.QtCore import Qt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from qfluentwidgets import FluentIcon as FIF

class RichEdit(QWidget):
    """ Rich edit """

    def __init__(self,image_path, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.layout = QVBoxLayout(self)
        self.setFixedWidth(350)
        # 创建 TextEdit（富文本编辑器）
        self.textEdit = TextEdit(self)

        # 创建工具栏
        self.create_toolbar(image_path)


        # 添加控件到布局
        self.layout.addWidget(self.textEdit)

        # 连接 QTextEdit 的选择变化信号
        # self.textEdit.selectionChanged.connect(self.update_font_size_combo)

    def create_toolbar(self,image_path):
        """创建工具栏并添加功能按钮"""
        self.toolbar = CommandBar(self)

        # 加粗
        bold_action = Action(FIF.FONT_INCREASE, self.tr('加粗'))
        bold_action.triggered.connect(self.set_bold)
        self.toolbar.addAction(bold_action)

        # 斜体
        italic_action = Action(FIF.PENCIL_INK, self.tr('斜体'))
        italic_action.triggered.connect(self.set_italic)
        self.toolbar.addAction(italic_action)

        # 下划线
        underline_action = Action(FIF.REMOVE, self.tr('下划线'))
        underline_action.triggered.connect(self.set_underline)
        self.toolbar.addAction(underline_action)

        self.toolbar.addSeparator()
        
        if image_path:
            watch_action = Action(FIF.ZOOM_IN, self.tr('放大'))
            watch_action.triggered.connect(lambda: self.parent.view_image(image_path))
            self.toolbar.addAction(watch_action)

        delete_action = Action(FIF.DELETE, self.tr('删除'))
        delete_action.triggered.connect(self.parent.delete_group)
        self.toolbar.addAction(delete_action)



        # 分点列表
        # bullet_action = Action(FIF.CALENDAR, self.tr('分点列表'))
        # bullet_action.triggered.connect(self.add_bullet_list)
        # self.toolbar.addAction(bullet_action)

        # # 字号选择框
        # self.font_size_combo = ComboBox(self)
        # self.font_size_combo.addItems([str(i) for i in range(8, 73, 2)])  # 字号从8到72（步长为2）
        # self.font_size_combo.currentTextChanged.connect(self.change_font_size)
        # self.toolbar.addWidget(self.font_size_combo)

        # 左对齐
        # align_left_action = Action(FIF.ALIGNMENT,self.tr('斜体'))
        # align_left_action.triggered.connect(self.align_left)
        # self.toolbar.addAction(align_left_action)

        # # 居中对齐
        # align_center_action = Action(FIF.MENU,self.tr('斜体'))
        # align_center_action.triggered.connect(self.align_center)
        # self.toolbar.addAction(align_center_action)

        # # 右对齐
        # align_right_action = Action(FIF.MENU, self.tr('斜体'))
        # align_right_action.triggered.connect(self.align_right)
        # self.toolbar.addAction(align_right_action)

        # 添加工具栏到布局
        self.layout.addWidget(self.toolbar)

    def set_bold(self):
        """设置选中文字为加粗"""
        fmt = self.textEdit.currentCharFormat()
        fmt.setFontWeight(QFont.Bold if fmt.fontWeight() != QFont.Bold else QFont.Normal)
        self.textEdit.mergeCurrentCharFormat(fmt)

    def set_italic(self):
        """设置选中文字为斜体"""
        fmt = self.textEdit.currentCharFormat()
        fmt.setFontItalic(not fmt.fontItalic())
        self.textEdit.mergeCurrentCharFormat(fmt)

    def set_underline(self):
        """设置选中文字为下划线"""
        fmt = self.textEdit.currentCharFormat()
        fmt.setFontUnderline(not fmt.fontUnderline())
        self.textEdit.mergeCurrentCharFormat(fmt)

    def add_bullet_list(self):
        """插入分点列表"""
        cursor = self.textEdit.textCursor()
        list_format = QTextListFormat()
        list_format.setStyle(QTextListFormat.ListDisc)
        cursor.createList(list_format)

    def change_font_size(self, size):
        """更改选中文字的字体大小"""
        fmt = self.textEdit.currentCharFormat()
        fmt.setFontPointSize(int(size))  # 设置字号
        self.textEdit.mergeCurrentCharFormat(fmt)

    # def update_font_size_combo(self):
    #     """更新字体大小选择框为选中文本的字号"""
    #     cursor = self.textEdit.textCursor()
    #     if cursor.hasSelection():
    #         fmt = cursor.charFormat()
    #         font_size = fmt.fontPointSize()
    #         # 更新字号选择框中的值
    #         if font_size:
    #             self.font_size_combo.setCurrentText(str(int(font_size)))
    #     else:
    #         self.font_size_combo.setCurrentText("")  # 清空字号框

    def align_left(self):
        """设置选中文本左对齐"""
        cursor = self.textEdit.textCursor()
        block_format = cursor.blockFormat()
        block_format.setAlignment(Qt.AlignmentFlag.AlignLeft)
        cursor.mergeBlockFormat(block_format)

    def align_center(self):
        """设置选中文本居中对齐"""
        cursor = self.textEdit.textCursor()
        block_format = cursor.blockFormat()
        block_format.setAlignment(Qt.AlignmentFlag.AlignCenter)
        cursor.mergeBlockFormat(block_format)

    def align_right(self):
        """设置选中文本右对齐"""
        cursor = self.textEdit.textCursor()
        block_format = cursor.blockFormat()
        block_format.setAlignment(Qt.AlignmentFlag.AlignRight)
        cursor.mergeBlockFormat(block_format)

    def save_to_word(self):
        """保存富文本内容到 Word 文档"""
        # 从 QTextEdit 获取 QTextDocument
        document = self.textEdit.document()
        # 初始化 Word 文档
        doc = Document()

        # 遍历 QTextDocument 的块（段落）
        block = document.firstBlock()
        while block.isValid():
            paragraph = doc.add_paragraph()  # 添加段落

            # 设置段落对齐方式
            block_format = block.blockFormat()
            alignment = block_format.alignment()
            if alignment == Qt.AlignmentFlag.AlignLeft:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif alignment in (Qt.AlignHCenter, Qt.AlignmentFlag.AlignCenter):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == Qt.AlignmentFlag.AlignRight:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 默认左对齐

            # 遍历段落中的文本片段
            it = block.begin()
            while not it.atEnd():
                fragment = it.fragment()
                if fragment.isValid():
                    char_format = fragment.charFormat()
                    text = fragment.text()

                    # 添加到 Word 段落的 Run
                    run = paragraph.add_run(text)

                    # 应用样式
                    if char_format.fontWeight() == QFont.Bold:
                        run.bold = True
                    if char_format.fontItalic():
                        run.italic = True
                    if char_format.fontUnderline():
                        run.underline = True
                    font_size = char_format.fontPointSize()
                    if font_size > 0:
                        run.font.size = font_size

                it += 1  # 进入下一个 fragment

            block = block.next()  # 移动到下一个块

