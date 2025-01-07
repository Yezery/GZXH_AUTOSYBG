import sys
from PyQt5.QtWidgets import QWidget,QVBoxLayout,QHBoxLayout,QDialog,QFileDialog,QGraphicsOpacityEffect,QSizePolicy
from PyQt5.QtCore import QEasingCurve,Qt,QStandardPaths,QPoint,pyqtSlot
from PyQt5.QtGui import QFont
from qfluentwidgets import ToolTipPosition,ToolTipFilter,ImageLabel,PrimaryToolButton,PrimaryPushButton,CardWidget,FlowLayout,BodyLabel,PushButton,LineEdit,TeachingTip,InfoBarIcon,TeachingTipTailPosition,pyqtSignal,CommandBarView,Action,FlyoutAnimationType,Flyout,StateToolTip,SmoothScrollArea
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
import os
from components.RichEdit import RichEdit
from components.Message import createMessage
from common.config import cfg
from qfluentwidgets import FluentIcon as FIF
from utils.ConvertFile import ConvertFileWorker
from view.summary_interface import SummaryInterface
from docx.document import Document as DocumentObject
from collections import OrderedDict
class ImageInputGroup(QWidget):
    def __init__(self, image_path, parent=None):
        super().__init__(parent)
        self.parentObject = parent
        self.setAcceptDrops(True)
        # 创建垂直布局
        self.layout = QVBoxLayout(self)
        self.adjustSize()

        if image_path:
            # 图片标签
            self.image_label = ImageLabel(image_path)
            self.image_label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            self.image_label.scaledToWidth(350)

        # 文件名编辑框
        self.filename_edit = RichEdit(image_path,self)
        # if image_path:
        #     self.filename_edit.textEdit.setText(os.path.basename(image_path).split('.')[0])
        self.filename_edit.textEdit.setFixedHeight(100)  # 设置固定高度
        self.filename_edit.textEdit.setPlaceholderText("请输入描述")

        # 创建添加按钮
        self.insert_group_button = QWidget(self)
        self.insert_group_button_layout = QHBoxLayout(self.insert_group_button)

        self.add_button = PrimaryPushButton(FIF.ADD,"插入图片", self)
        self.add_button.setAcceptDrops(True)
        self.add_text_button = PrimaryPushButton(FIF.ADD,"插入纯文字", self)
        self.insert_group_button.setVisible(False)  # 初始时不可见
        self.insert_group_button_layout.addWidget(self.add_button,1)
        self.insert_group_button_layout.addWidget(self.add_text_button,1)
        self.add_text_button.clicked.connect(lambda:self.parentObject.insert_before_key(ImageInputGroup(None,self.parentObject),None,self))
        self.add_button.clicked.connect(self.click_upload)
        # 将按钮添加到布局中
        self.layout.addWidget(self.insert_group_button)

        # 创建水平布局，用于将按钮布局放在输入框右侧
        input_layout = QVBoxLayout()
        # input_layout.addLayout(button_layout)  # 左边为按钮垂直布局
        input_layout.addWidget(self.filename_edit)  # 右边为输入框

        # 将图片标签和输入框布局加入主布局
        if image_path:
            self.layout.addWidget(self.image_label)
        self.layout.addLayout(input_layout)

        # 用于存储输入框与图片路径的映射
        if image_path:
            self.image_path = image_path
        else:
            self.image_path = None
        self.update_image_map()
        self.filename_edit.textEdit.textChanged.connect(self.update_image_map)

    def click_upload(self):
        file_dialog = QFileDialog(self)
        file_dialog.setFileMode(QFileDialog.ExistingFiles)
        file_dialog.setNameFilters(["Image files (*.png *.jpg)"])
        if file_dialog.exec_():
            file_paths = file_dialog.selectedFiles()
            self.parentObject.insert_before_key(ImageInputGroup(file_paths[0],self.parentObject),file_paths[0],self)
        

    def update_image_map(self):
        """更新输入框对象与图片路径的映射"""
        parent = self.parentObject
        if parent:
            parent.input_image_map[self] = self.image_path
            
    def delete_group(self):
        """删除图片输入组"""
        parent = self.parentObject  # 获取父组件
        if parent is None:
            return

        # 从父布局中移除控件
        layout = parent.layout
        if layout.indexOf(self) >= 0:
            layout.removeWidget(self)

        # 从映射字典中移除该控件
        if self in parent.input_image_map:
            del parent.input_image_map[self]

        # 销毁控件资源
        self.setParent(None)  # 将控件从父布局中分离
        self.deleteLater()    # 延迟销毁，确保资源释放
        
        parent.visible()
        # 更新父布局
        parent.layout.update()

    def view_image(self, image_path):
        """放大查看图片"""
        image_label = ImageLabel(image_path)
        image_label.scaledToWidth(900)
        zoomed_window = QDialog(self)
        zoomed_window.setWindowTitle("放大查看")
        zoomed_layout = QVBoxLayout(zoomed_window)

        # 缩放图像并保持平滑
        zoomed_layout.addWidget(image_label)
        zoomed_window.exec_()

    def enterEvent(self, event):
        """鼠标进入时显示按钮"""
        self.insert_group_button.setVisible(True)

    def leaveEvent(self, event):
        """鼠标离开时隐藏按钮"""
        self.insert_group_button.setVisible(False)

    # # 拖拽事件处理
    # def dragEnterEvent(self, event):
    #     """处理拖拽进入"""
    #     if event.mimeData().hasUrls():  # 确保拖拽的是文件
    #         event.accept()
    #     else:
    #         event.ignore()

    # def dropEvent(self, event):
    #     """处理文件放下事件"""
    #     if event.mimeData().hasUrls():
    #         urls = event.mimeData().urls()  # 获取拖拽的文件路径
    #         file_path = urls[0].toLocalFile()  # 获取第一个文件的路径
    #         if file_path.lower().endswith(('png', 'jpg', 'jpeg')):  # 只处理图片
    #             self.parentObject.insert_before_key(ImageInputGroup(file_path, self.parentObject), file_path, self)

class DropFileUploadImages(CardWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parentObject = parent
        self.input_image_map = {}
        # 设置窗口属性
        self.setAcceptDrops(True)
        self.setBorderRadius(10)
        self.setStyleSheet("QWidget{background:transparent;border:none;}")
        # 主布局 - 使用 FlowLayout 或其他布局（根据需求）
        self.layout = FlowLayout(self, needAni=True)
        self.layout.setAnimation(250, QEasingCurve.OutQuad)

        # 标签提示
        self.label = BodyLabel("拖拽文件 或 点击选择图片文件", self)
        self.label.setAlignment(Qt.AlignCenter)  # 居中对齐文本
        self.label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        # 设置占位符的样式表
        self.label.setStyleSheet(
            """
            QLabel {
                border-radius: 8px;        /* 圆角 */
                font-size: 14px;
                color: #555;
                background-color: rgba(255, 255, 255, 0); /* 背景色与透明度 */
            }
            QLabel:hover {
                color: #0078d4;
                border: 2px dashed #0078d4; /* 虚线边框 */
            }
            """
        )
        self.visible()

    def resizeEvent(self, event):
        """确保容器控件和子控件随父控件大小变化而更新位置和大小"""
        super().resizeEvent(event)
        # 重新定位按钮组和标签
        self.label.setGeometry(0, 0, self.width(), self.height())  # 标签填满容器

    def dragEnterEvent(self, event):
        """鼠标拖入事件"""
        if event.mimeData().hasUrls():
            event.accept()
            self.label.setText("释放文件到此区域")
            
            # 获取当前拖动位置
            drag_position = event.pos()
            
            # 遍历布局中的控件，检查是否鼠标进入了某个 ImageInputGroup
            for child in self.children():
                if isinstance(child, ImageInputGroup):
                    if child.rect().contains(child.mapFromParent(drag_position)):
                        # 触发该控件的 enterEvent
                        child.enterEvent(event)
                        break
        else:
            event.ignore()

    def dragLeaveEvent(self, event):
        """鼠标拖出事件"""
        self.label.setText("拖拽文件 或 点击选择图片文件")

    def dropEvent(self, event):
        """鼠标放开事件"""
        if event.mimeData().hasUrls():
            file_paths = [url.toLocalFile() for url in event.mimeData().urls()]
            self.add_image_groups(file_paths)
        else:
            self.label.setText("不支持的文件类型")
        if not self.input_image_map:
            self.label.setVisible(True)
        else:
            self.label.setVisible(False)
        self.label.setText("拖拽文件 或 点击选择图片文件")

    def add_image_groups(self, file_paths):
        """根据上传的图片文件路径动态添加图片组"""
        for path in file_paths:
            if path.lower().endswith(('.png', '.jpg')):
                image_group = ImageInputGroup(path, self)
                self.layout.addWidget(image_group)
        self.visible()

    def visible(self):
        if  self.input_image_map == {}:
            self.parentObject.tools_button_group.clear_button.setVisible(False)
            self.layout.update()
            self.label.setVisible(True)
        else:
            self.parentObject.tools_button_group.clear_button.setVisible(True)
            self.layout.update()
            self.label.setVisible(False)

    def clear_all_image_groups(self):
        """清除所有图片组"""
        # 遍历映射字典的所有键（控件），依次删除
        for image_group in list(self.input_image_map.keys()):
            image_group.delete_group()
        # 清空映射字典
        self.input_image_map.clear()

    def add_no_image_description(self):
        """添加无图描述"""
        no_image_group = ImageInputGroup(None,self)
        self.layout.addWidget(no_image_group)
        self.visible()

    def mousePressEvent(self, event):
        """鼠标点击事件，支持点击上传文件"""
        if event.button() == Qt.LeftButton:
            file_dialog = QFileDialog(self)
            file_dialog.setFileMode(QFileDialog.ExistingFiles)
            file_dialog.setNameFilters(["Image files (*.png *.jpg)"])
            if file_dialog.exec_():
                file_paths = file_dialog.selectedFiles()
                self.add_image_groups(file_paths)
        if not self.input_image_map:
            self.label.setVisible(True)
        else:
            self.label.setVisible(False)

    def insert_before_key(self,new_key, new_value, target_key):
        # 将普通字典临时转换为 OrderedDict
        ordered_map = OrderedDict(self.input_image_map)

        # 获取所有键的顺序
        keys = list(ordered_map.keys())
        
        # 确保目标键存在
        if target_key not in ordered_map:
            raise KeyError(f"Key '{target_key}' not found in the dictionary.")
        
        # 获取目标键的位置
        index = keys.index(target_key)
        
        # 使用 move_to_end() 来插入新元素
        for key in keys[:index]:
            ordered_map.move_to_end(key)  # 将目标键前的所有键移到末尾

        ordered_map[new_key] = new_value  # 插入新的键值对
        ordered_map.move_to_end(new_key, last=False)  # 确保新插入的键在目标键前

        # 打印修改后的 OrderedDict
        print(ordered_map)

        # 将 OrderedDict 转换回普通字典并赋值给 self.input_image_map
        self.input_image_map = dict(ordered_map)
        # 找到目标控件在布局中的索引位置
        target_widget = self.layout.itemAt(index)  # 获取目标键的当前布局项
        if target_widget:
            target_index = self.layout.indexOf(target_widget)
            # 使用 insertWidget 来插入到指定位置
            self.layout.insertWidget(target_index, new_key)
        else:
            # 如果目标控件未找到，直接将新的控件添加到末尾
            self.layout.addWidget(new_key)
        self.visible()
class ToolsButtonGroup(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parentObject = parent
       # 按钮组
        self.button_group_widget = QWidget(self)
        self.button_group_layout = QHBoxLayout(self.button_group_widget)
        self.button_group_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        self.button_group_layout.setContentsMargins(0, 0, 0, 0)
        # 清除全部按钮
        self.clear_button = PrimaryPushButton(FIF.DELETE,"清除全部", self)
        self.clear_button.setFixedWidth(165)
        self.clear_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
       

        # 添加无图描述按钮
        self.add_no_image_button = PrimaryPushButton(FIF.ADD,"添加无图描述", self)
        self.add_no_image_button.setFixedWidth(165)
        self.add_no_image_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
       

        # 将按钮添加到按钮布局中
        self.button_group_layout.addWidget(self.add_no_image_button)
        self.button_group_layout.addWidget(self.clear_button)
        self.layout = QHBoxLayout(self)
        self.layout.addWidget(self.button_group_widget)

class FileButtonGroup(QWidget):
    def __init__(self, file_name,file,file_type ,parent=None):
        super().__init__(parent)
        self.parentObject =parent
        self.file_name = file_name
        self.file = file
        self.file_type =  file_type
        # 按钮的布局
        layout = QVBoxLayout(self)

        # 创建按钮
        self.button = PushButton(self.file_type, self)
        self.button.setFixedSize(100, 100)  # 设置正方形按钮
        # self.button.setStyleSheet('font: 12px "Microsoft YaHei"; border-radius: 10px;border: 1px solid #ccc;')

        # 连接按钮点击事件
        self.button.clicked.connect(self.createCommandBarFlyout)

        layout.addWidget(self.button)

        # 设置布局
        self.setLayout(layout)

    def createCommandBarFlyout(self):
        TeachingTip.create(
                target=self,
                icon=InfoBarIcon.INFORMATION,
                title='文件名',
                content=self.file_name,
                isClosable=True,
                tailPosition=TeachingTipTailPosition.BOTTOM,
                duration=2000,
                parent=self
            )   

        """弹出操作工具条（保存和删除）"""
        view = CommandBarView(self)

        # 创建动作
        view.addAction(Action(FIF.SAVE, self.tr('Save'), triggered=self.saveFile))
        view.addAction(Action(FIF.DELETE, self.tr('Delete'), triggered=self.delete_button_group))
        view.resizeToSuitableWidth()

        x = self.width()  # 获取当前组件的宽度
        pos = self.mapToGlobal(QPoint(x-50, -5))
        Flyout.make(view, pos, self, FlyoutAnimationType.FADE_IN)

    def saveFile(self):
        save_path = QFileDialog.getExistingDirectory(self,directory=QStandardPaths.writableLocation(QStandardPaths.DesktopLocation),caption= "选择保存路径")
        parent = self.parent().parent().parent().parent().parent().parent().findChild(AppInterface)
        if save_path == "":
            createMessage(parent.right_scroll_area,"警告", "用户取消",2)
        else:
            match self.file_type:
                case "docx":
                    self.file.save(os.path.join(save_path, self.file_name))
                case "pdf":
                    with open(f'{os.path.join(save_path, self.file_name)}', 'wb') as output_pdf:
                        output_pdf.write(self.file.getvalue())
            createMessage(parent.right_scroll_area,"成功", "保存成功",1)

    def delete_button_group(self):
        """删除按钮组"""
        parent = self.parent()
        if parent:
            parent.remove_result_file(self)  # 请求父级删除该按钮组
        self.deleteLater()  # 销毁当前控件
        parent.file_buttons_map.pop(self.file_name)

class ResultFile(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parentObject = parent

        # 设置布局
        self.flowlayout = FlowLayout(self, needAni=True)

        # Customize animation
        self.flowlayout.setAnimation(250, QEasingCurve.OutQuad)

        # 初始化文件按钮列表，用于管理已添加的按钮
        self.file_buttons_map = {}

        # 添加一些示例文件按钮
        # self.add_result_file("aiko.pdf", self.flowlayout)

        self.setLayout(self.flowlayout)
        self.resize(300, 400)
        self.setStyleSheet('QPushButton{padding: 5px 10px; font: 15px "Microsoft YaHei"}')

    def remove_result_file(self, file_button_group):
        """根据按钮组移除文件"""
        layout = self.layout()  # 获取当前布局

        # 从布局中移除 widget
        layout.removeWidget(file_button_group)
        # 从文件按钮字典中删除
        for key, value in list(self.file_buttons_map.items()):
            if value == file_button_group:
                del self.file_buttons_map[key]
                break
        
        # 销毁按钮组
        file_button_group.deleteLater()
        layout.update()
            
class ProcessTask(CardWidget):
    def __init__(self, text, parent=None):
        super().__init__(parent)
        self.parentObject = parent
        self.setObjectName(text.replace(' ', '-'))
        self.layout = QVBoxLayout(self)
        self.doc = None
        self.layout.setAlignment(Qt.AlignCenter)  # 确保布局内控件居中
        # 输入框
        self.docx_rename = LineEdit(self)
        self.docx_rename.installEventFilter(ToolTipFilter(self.docx_rename,showDelay=10,position=ToolTipPosition.BOTTOM))
        self.docx_rename.setToolTip(self.tr('只需保留实验报告名称，不用加个人信息哦'))
        # 获取输入框内容
        self.docx_rename.setPlaceholderText("重命名输出文件名")
        self.buttonGroupLayout = QHBoxLayout(self)
 
        self.toDOCX = PrimaryPushButton("输出为DOCX")
        self.toDOCX.setEnabled(False)
        self.toPDF = PrimaryPushButton("输出为PDF")
        self.toPDF.setEnabled(False)
        self.layout.addWidget(self.docx_rename)
        self.layout.addLayout(self.buttonGroupLayout)
        self.buttonGroupLayout.addWidget(self.toDOCX)
        self.buttonGroupLayout.addWidget(self.toPDF)
   
        self.toDOCX.clicked.connect(self.toDOCXTask)
        self.toPDF.clicked.connect(self.toPDFTask)

        self.setBorderRadius(10)
        
    def mainTask(self):
        images = []  # 用于保存图片路径
        texts = []  # 用于保存输入框中的文字
        # 遍历 input_image_map，获取每个输入框的文字和图片路径
        for image_group, image_path in self.parentObject.drop_upload.input_image_map.items():
            text = image_group.filename_edit.textEdit.document()
            # 根据输入框和图片路径的情况，分别处理
            if text and image_path:
                # 如果有文字并且有图片
                texts.append(text)
                images.append(image_path)
            elif text and not image_path:
                # 如果有文字但没有图片
                texts.append(text)
                images.append(None)  # 没有图片，则将其设置为None
            elif not text and image_path:
                # 如果没有文字但有图片
                texts.append("无文字说明")  # 可以自定义文字说明，比如设置为"无文字说明"
                images.append(image_path)
            else:
                # 既没有文字也没有图片的情况
                texts.append("无文字说明")
                images.append(None)
        header = f"实验人：{cfg.get(cfg.userName)}（学号：{cfg.get(cfg.userId)}）"
        # 修改页眉
        self.modify_header(header)
        # 文档清洗
        self.doc_clean()
        # 插入实验过程
        self.insert_paragraphs(images, texts)
        # 插入心得
        self.insert_learning()
     # 插入实验过程
    def insert_paragraphs(self, images, texts):
        """插入实验过程"""
        pic_num = 1
        for i, (image, text) in enumerate(zip(images, texts)):
            if text != "无文字说明":
                # 遍历 QTextDocument 的块（段落）
                block = text.firstBlock()
                while block.isValid():
                    text_paragraph = self.doc.add_paragraph()  # 添加段落
                    # 设置段落对齐方式
                    block_format = block.blockFormat()
                    alignment = block_format.alignment()
                    if alignment == Qt.AlignLeft:
                        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif alignment in (Qt.AlignHCenter, Qt.AlignCenter):
                        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif alignment == Qt.AlignRight:
                        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 默认左对齐

                    # 遍历段落中的文本片段
                    it = block.begin()
                    while not it.atEnd():
                        fragment = it.fragment()
                        if fragment.isValid():
                            char_format = fragment.charFormat()
                            text = fragment.text()

                            # 添加到 Word 段落的 Run
                            run = text_paragraph.add_run(text)
                            if image:
                                text_paragraph.add_run(f"，如图 {pic_num}所示")
                            # 应用样式
                            if char_format.fontWeight() == QFont.Bold:
                                run.bold = True
                            if char_format.fontItalic():
                                run.italic = True
                            if char_format.fontUnderline():
                                run.underline = True
                            font_size = char_format.fontPointSize()
                            if font_size > 0:
                                run.font.size = font_size
                            # 在图片上方插入文字
                            text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 水平居左
                            # 设置段落间距为 1.5 倍行距
                            text_paragraph.paragraph_format.line_spacing = Pt(
                                24
                            )  # 1.5倍行距（16磅*1.5=24磅）
                            run.font.size = Pt(10.5)  # 5号字（10.5磅）
                            run.font.name = "宋体"  # 设置字体为宋体
                            spacing_paragraph = self.doc.add_paragraph()  # 使用文档对象添加间隙段落
                            spacing_paragraph.paragraph_format.space_after = Pt(
                                1)
                        it += 1  # 进入下一个 fragment
                    block = block.next()  # 移动到下一个块
                    
            if image:
                # 插入图片
                image_paragraph = self.doc.add_paragraph()  # 使用文档对象添加新段落
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_paragraph.add_run().add_picture(image, width=Inches(5))  # 调整宽度
                # 添加图片题注（段落）
                caption_paragraph = self.doc.add_paragraph(f"图 {pic_num}")
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # spacing_paragraph = self.doc.add_paragraph()  # 使用文档对象添加间隙段落
                # spacing_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # spacing_paragraph.paragraph_format.space_after = Pt(
                #     12)
                pic_num += 1
    def get_app_path(self):
        """
        获取应用程序运行目录：
        - 在未打包时，返回当前脚本所在目录
        - 在打包时，根据平台获取资源路径
        """
        if getattr(sys, "frozen", False):  # 检测是否为打包环境
            if sys.platform == "darwin":  # macOS 的 .app 文件
                return os.path.join(os.path.dirname(sys.executable), "../Resources")
            else:  # Windows/Linux 的打包目录
                return os.path.dirname(sys.executable)
        else:
            # 未打包时，返回脚本所在目录
            return os.path.abspath(os.path.dirname(__file__))
        
    def toPDFTask(self):
        self.stateTooltip = None
        if not self.stateTooltip:
                self.stateTooltip = StateToolTip('正在输出中', '请耐心等待哦～', self.parentObject)
                self.stateTooltip.move(self.parentObject.geometry().topRight() - QPoint(270, -10))
                self.stateTooltip.show()
        self.toPDF.setEnabled(False)
        # formatPath = f"{Config['SAVE_PATH']}/{Config['COURSE']}_{Config['ID']}_{Config['NAME']}_{self.docx_rename.text()}"
        self.mainTask()
        self.doc.save(os.path.join(self.get_app_path(), "temp.docx"))
        """启动文件转换线程，处理后台转换任务"""
        self.worker = ConvertFileWorker(f"{self.get_app_path()}/temp.docx","",2)

        # 连接信号，使用类方法作为回调
        self.worker.finished.connect(self.on_conversion_finished)
        self.worker.error.connect(self.on_conversion_error)

        # 启动线程
        self.worker.start()

    def on_conversion_finished(self, file):
        self.add_result_list(file, "pdf")
        """文件转换完成时的回调函数"""
        self.worker.deleteLater()  # 清理线程资源
        if self.stateTooltip:
            self.stateTooltip.setContent('输出完成啦 😆')
            self.stateTooltip.setState(True)
            self.stateTooltip = None
        self.toPDF.setEnabled(True)
        # 在这里执行其他操作，例如更新 UI 或通知用户

    def on_conversion_error(self, error_message):
        """文件转换失败时的回调函数"""
        print(f"转换失败：{error_message}")
        self.worker.deleteLater()  # 清理线程资源
        if self.stateTooltip:
            self.stateTooltip.setContent('输出失败了 😭')
            self.stateTooltip.setState(True)
            self.stateTooltip = None
        self.toPDF.setEnabled(True)
        # 在这里执行错误处理操作，例如提示用户转换失败
        
    def toDOCXTask(self):
        if not self.doc:
            return
        try:
            self.mainTask()
            self.add_result_list(self.doc, "docx")
        except Exception as e:
            createMessage(self.parentObject, title="发生错误", message=f"{e}", _type=0)

    def add_result_list(self,file,_type="docx"):
        result_file_object = self.parentWidget().children()[-1].findChild(ResultFile)
        formatName = f"{cfg.get(cfg.userCourse)}_{cfg.get(cfg.userId)}_{cfg.get(cfg.userName)}_{self.docx_rename.text()}.{_type}"
        if result_file_object.file_buttons_map.get(formatName):
            result_file_object.file_buttons_map[formatName] = FileButtonGroup(formatName,file,_type ,self)
        else:
            result_file_object.file_buttons_map[formatName] = FileButtonGroup(formatName,file,_type,self)
            result_file_object.flowlayout.addWidget(FileButtonGroup(formatName,file,_type ,self))
        createMessage(self.parentObject,"输出", "输出成功,请点击文件列表进行保存！",1)
        
    # 修改页眉
    def modify_header(self, new_header_text):
        # 修改每个节的页眉
        for section in self.doc.sections:
            header = section.header
            # 清空原有内容
            header.paragraphs[0].clear()
            # 添加新的内容
            header.paragraphs[0].add_run(new_header_text)

    # 插入心得
    def insert_learning(self):
        # 添加心得栏
        text_paragraph = self.doc.add_paragraph()
        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 水平居左
        # 设置段落间距为 1.5 倍行距
        text_paragraph.paragraph_format.line_spacing = Pt(
            24
        )  # 1.5倍行距（16磅*1.5=24磅）
        run = text_paragraph.add_run("【实验总结（个人心得）】")
        run.font.size = Pt(10.5)  # 5号字（10.5磅）
        run.font.name = "宋体"  # 设置字体为宋体
        run.font.bold = True

        def format_text(content):
            paragraphs = content.split("\n\n")  # 将内容按段落分开
            for paragraph_text in paragraphs:
                text_paragraph = self.doc.add_paragraph()  # 使用文档对象添加文字段落
                text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 水平居左

                # 设置段落间距为 1.5 倍行距
                text_paragraph.paragraph_format.line_spacing = Pt(
                    24
                )  # 1.5倍行距（16磅*1.5=24磅）

                # 设置段落开头缩进2个字（24磅，约等于2个汉字宽度）
                text_paragraph.paragraph_format.first_line_indent = Pt(24)

                # 添加段落内容
                run = text_paragraph.add_run(paragraph_text)
                run.font.size = Pt(10.5)  # 设置字体为5号字（10.5磅）
                run.font.name = "宋体"  # 设置字体为宋体

        format_text(self.parentObject.summary_widget.summary_text.toPlainText())

    # 文档清洗
    def doc_clean(self):
        # 标记是否找到 "【实验过程记录】"
        found = False
        # 遍历文档中的段落
        for para in self.doc.paragraphs:
            if found:
                CT_P = para._element
                CT_P.getparent().remove(CT_P)
            elif "【实验过程记录】" in para.text or "【实验步骤】" in para.text:
                # 找到 "【实验过程记录】" 后标记，跳过该段落
                found = True
                # 保留该段落，不清除其内容
                continue

class AppInterface(QWidget):
    
    def __init__(self, text: str, parent=None):
        super().__init__(parent=parent)
        self.setObjectName(text.replace(' ', '-'))
        self.__initWidget()
        self.__initLayout()
        self.summary_widget.isUpload.connect(self.set_upload)
        self.summary_widget.docx_emit.connect(self.set_doc)

    def __initWidget(self):
        # 创建主布局
        self.main_layout = QHBoxLayout(self)        
        # 在左侧滚动区域中放置一个 QWidget 作为容器
        self.left_widget = QWidget(self)
        self.left_layout = QVBoxLayout(self.left_widget)
        self.summary_widget = SummaryInterface("app summary interface", self)
        self.process_task = ProcessTask("app process task", self)
        self.left_result_file_scroll_area = SmoothScrollArea(self)
        self.result_file = ResultFile(self) 

        self.left_layout.addWidget(self.summary_widget)
        self.left_layout.addWidget(self.process_task)
        self.left_layout.addWidget(self.left_result_file_scroll_area)
        self.left_result_file_scroll_area.setWidget(self.result_file)

        # 创建右侧滚动区域
        self.tools_button_group = ToolsButtonGroup(self)
        self.drop_upload = DropFileUploadImages(self)
        self.right_widget = QWidget(self)
        self.right_widget.setContentsMargins(0,0, 0, 0)
        self.right_layout = QVBoxLayout(self.right_widget)        

        self.right_scroll_area = SmoothScrollArea(self)

        self.right_layout.addWidget(self.right_scroll_area)
        self.right_layout.addWidget(self.tools_button_group)
        self.right_scroll_area.setWidget(self.drop_upload)
        
        # 将左侧和右侧区域添加到主布局
        self.main_layout.addWidget(self.left_widget,1) 
        self.main_layout.addWidget(self.right_widget,1)

    def __initLayout(self):

        self.left_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        self.left_widget.setStyleSheet("background:transparent;border:none;")
        self.left_layout.setAlignment(Qt.AlignTop)
        self.left_layout.setContentsMargins(0,0, 0, 0)

        self.right_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        self.right_widget.setStyleSheet("background:transparent;border:none;")
        self.right_layout.setAlignment(Qt.AlignTop)
        self.right_layout.setContentsMargins(0,0, 0, 0)
        
        self.summary_widget.setBorderRadius(10)
        self.summary_widget.setFixedHeight(250)

        self.left_result_file_scroll_area.setAutoFillBackground(True)
        self.left_result_file_scroll_area.enableTransparentBackground()
        self.left_result_file_scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.left_result_file_scroll_area.setViewportMargins(0, 0, 0, 0)
        self.left_result_file_scroll_area.setWidgetResizable(True)

        self.right_scroll_area.enableTransparentBackground()
        self.right_scroll_area.setWidgetResizable(True)

        self.result_file.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)

        self.tools_button_group.clear_button.clicked.connect(self.drop_upload.clear_all_image_groups)
        self.tools_button_group.add_no_image_button.clicked.connect(self.drop_upload.add_no_image_description)
        

    @pyqtSlot(DocumentObject)
    def set_doc(self,docx):
        self.process_task.doc = docx
        self.process_task.toDOCX.setEnabled(True)
        self.process_task.toPDF.setEnabled(True)

    @pyqtSlot(bool)
    def set_upload(self,r:bool):
        if not r:
            # 只要文件名称并去掉后缀
            self.process_task.docx_rename.setText(os.path.splitext(os.path.basename(self.summary_widget.upload_input.file_paths[0]))[0])
