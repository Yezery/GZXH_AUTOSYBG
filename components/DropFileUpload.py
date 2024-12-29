import os
from PyQt5.QtWidgets import QWidget,QVBoxLayout,QFileDialog
from PyQt5.QtCore import QEvent,Qt,QPoint,pyqtSlot
from PyQt5.QtGui import QDragEnterEvent
from docx import Document
from docx.document import Document as DocumentObject
from components.Message import createMessage
from task.LoadConfig import get_config
from task import AI
from qfluentwidgets import pyqtSignal,BodyLabel,StateToolTip
from qfluentwidgets import FluentIcon as FIF

class DropFileUploadDOCX(QWidget):
    ai_summary_generated = pyqtSignal(str) 
    isUpload = pyqtSignal(bool) 
    docx_emit = pyqtSignal(DocumentObject)
    start_process_signal = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.doc = None
        # 设置窗口属性
        self.setAcceptDrops(True)  # 启用拖拽功能
        self.stateTooltip = None
        # 主布局
        self.layout = QVBoxLayout(self)
        self.layout.setAlignment(Qt.AlignTop)

        self.keywords = ["【实验目的】", "【实验内容】", "【实验要求】"]

        # 标签显示文件路径
        self.label = BodyLabel("拖拽文件 / 点击选择实验报告模版", self)
        self.label.setFixedHeight(50)
        self.label.setStyleSheet(
            """
            QLabel {
                font-size: 12px;
                color: #555;
                border: 2px dashed #aaa;
                border-radius: 8px;
                padding: 10px;
                text-align: center;
            }
            QLabel:hover {
                border-color: #0078d4;
            }
            """
        )
        self.label.setAlignment(Qt.AlignCenter)
        self.layout.addWidget(self.label)

        # 鼠标点击事件
        self.label.mousePressEvent = self.openFileDialog

        self.start_process_signal.connect(self.request_ai)

    def dragEnterEvent(self, event: QDragEnterEvent):
        """鼠标拖入事件"""
        if event.mimeData().hasUrls():
            event.accept()
            self.label.setText("释放文件到此区域")
            self.label.setStyleSheet(self.label.styleSheet() + "background-color: #e8f4fc;")
        else:
            event.ignore()

    def dragLeaveEvent(self, event: QEvent):
        """鼠标拖出事件"""
        self.resetLabel()

    def dropEvent(self, event):
        """鼠标放开事件"""
        if event.mimeData().hasUrls():
            self.file_paths = [url.toLocalFile() for url in event.mimeData().urls()]
            if len(self.file_paths) > 1:  # 检查是否上传了多个文件
                createMessage(self.parent, title="警告", message="仅支持上传一个文件", _type=2)
                self.resetLabel()
                return
            if self.initDocx(self.file_paths):
                file_names = [os.path.basename(path) for path in self.file_paths]  # 只获取文件名
                self.label.setText(f"文件名称：".join(self.truncate_file_name(name) for name in file_names))
                self.label.setStyleSheet(self.label.styleSheet() + "background-color: #f9f9f9;")
                # self.request_ai()
            else:
                createMessage(self.parent, title="警告", message=f"文档中不符合要求", _type=2)
                self.label.setText("文档中不符合要求")
                
        else:
            self.label.setText("不支持的文件类型")

    def openFileDialog(self, event):
        """打开文件选择对话框"""
        self.file_paths, _ = QFileDialog.getOpenFileNames(self, "选择文件", filter="DOCX 文件 (*.docx);;所有文件 (*.*)")
        if self.file_paths:
            if self.initDocx(self.file_paths):
                file_names = [os.path.basename(path) for path in self.file_paths]  # 只获取文件名
                self.label.setText(f"文件名称：".join(self.truncate_file_name(name) for name in file_names))
                # self.request_ai()
            else:
                createMessage( self.parent,title="警告", message=f"文档中不符合要求", _type=2)
                self.label.setText("文档中不符合要求")
        else:
            self.resetLabel()

    def resetLabel(self):
        """重置标签内容"""
        self.label.setText("拖拽文件或点击选择文件")
        self.label.setStyleSheet(self.label.styleSheet() + "background-color: #f9f9f9;")

    def truncate_file_name(self, file_name, max_length=20):
        """截断文件名"""
        if len(file_name) > max_length:
            return file_name[:max_length] + "..."
        return file_name

    def initDocx(self, file_paths):
        try:
            """检查文档中是否包含所有关键字"""
            content = ""
            # 获取文档中所有段落的文本内容
            for para in Document(file_paths[0]).paragraphs:
                content += para.text
            # 检查关键字是否存在于文档内容中
            for keyword in self.keywords:
                if keyword not in content:
                    self.isUpload.emit(False)
                    return False
            self.doc = Document(file_paths[0])
            self.isUpload.emit(True) 
            self.docx_emit.emit(self.doc)
            return True
        except Exception as e:
            self.isUpload.emit(False)
            return False
        
    def request_ai(self):
        if not self.stateTooltip:
            self.stateTooltip = StateToolTip('正在生成中', '请耐心等待哦～', self.parent)
            self.stateTooltip.move(self.parent.geometry().topRight() - QPoint(270, -10))
            self.stateTooltip.show()
            Config = get_config().config
            if self.doc is not None:
                content = ""
                capture = False
                start_keyword = "【实验题目】"
                end_keyword = "【实验要求】"

                # 提取文档内容
                for para in self.doc.paragraphs:
                    if start_keyword in para.text:
                        capture = True
                        content += para.text + "\n"
                    elif capture:
                        content += para.text + "\n"
                        if end_keyword in para.text:
                            break
                # 启动后台线程处理
                self.worker = AI.QFAIWorker(
                    api_key=Config["API_KEY"],
                    secret_key=Config["SECRET_KEY"],
                    doc_content=content.strip(),
                )
                self.worker.finished.connect(self.on_ai_success)
                self.worker.error.connect(self.on_ai_error)
                self.worker.start()
           


    @pyqtSlot(str)
    def on_ai_success(self, ai_summary):
        """处理生成成功的信号"""
        self.ai_summary_generated.emit(ai_summary)
        if self.stateTooltip:
            self.stateTooltip.setContent('生成完成啦 😆')
            self.stateTooltip.setState(True)
            self.stateTooltip = None

    @pyqtSlot(str)
    def on_ai_error(self, error_message):
        """处理生成失败的信号"""
        self.ai_summary_generated.emit(error_message)
        if self.stateTooltip:
            self.stateTooltip.setContent(f'生成失败了 😭')
            self.stateTooltip.setState(True)
            self.stateTooltip = None